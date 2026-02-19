"""
Document processing service for extracting text from various file formats.
Supports PDF, DOCX, and TXT files.
"""
import logging
from pathlib import Path
from typing import Optional, Dict, Any
import pypdf
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing and extracting text from documents."""
    
    SUPPORTED_TYPES = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "text/plain": ".txt"
    }
    
    @staticmethod
    def is_supported(content_type: str) -> bool:
        """
        Check if the content type is supported.
        
        Args:
            content_type: MIME type of the file
            
        Returns:
            True if supported, False otherwise
        """
        return content_type in DocumentProcessor.SUPPORTED_TYPES
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            text_content = []
            metadata = {"page_count": 0}
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                metadata["page_count"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"[Page {page_num}]\n{text}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {e}")
                        continue
            
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                raise ValueError("No text could be extracted from PDF")
            
            logger.info(f"Extracted {len(full_text)} characters from {metadata['page_count']} pages")
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_texts.append(row_text)
            
            # Combine all text
            all_text = paragraphs + table_texts
            full_text = "\n\n".join(all_text)
            
            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables)
            }
            
            if not full_text.strip():
                raise ValueError("No text could be extracted from DOCX")
            
            logger.info(f"Extracted {len(full_text)} characters from DOCX with {metadata['paragraph_count']} paragraphs")
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from a TXT file.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text.strip():
                raise ValueError("Text file is empty")
            
            metadata = {
                "line_count": len(text.splitlines())
            }
            
            logger.info(f"Extracted {len(text)} characters from TXT file")
            return text, metadata
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                metadata = {
                    "line_count": len(text.splitlines()),
                    "encoding": "latin-1"
                }
                
                logger.info(f"Extracted {len(text)} characters from TXT file (latin-1 encoding)")
                return text, metadata
            except Exception as e:
                logger.error(f"Error processing TXT file with alternative encoding: {e}")
                raise
        except Exception as e:
            logger.error(f"Error processing TXT file {file_path}: {e}")
            raise
    
    @staticmethod
    def process_document(file_path: str, content_type: str) -> tuple[str, Dict[str, Any]]:
        """
        Process a document and extract text based on its content type.
        
        Args:
            file_path: Path to the document file
            content_type: MIME type of the document
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            ValueError: If content type is not supported
            Exception: If processing fails
        """
        if not DocumentProcessor.is_supported(content_type):
            raise ValueError(f"Unsupported content type: {content_type}")
        
        logger.info(f"Processing document: {file_path} (type: {content_type})")
        
        if content_type == "application/pdf":
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif content_type == "text/plain":
            return DocumentProcessor.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported content type: {content_type}")
